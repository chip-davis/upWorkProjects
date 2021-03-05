import os
import pandas as pd
import pyexcel
from pyexcel._compact import OrderedDict
import xlrd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, Inches
from docx2pdf import convert

cwd = os.getcwd()


def getDictionary():
    file    = 'data.xlsx'
    
    df = pd.read_excel(file, sheet_name='Sheet1')
    dic = (df.to_dict('index'))
    length = len(dic)
    count = 0
    while count < length:
        for key, value in dic.get(count).items():
            try:
                key = key.strip()
                value = value.strip()
            except:
                continue
        count += 1
    return dic, length



def getSubHeadings():
    with open("subheadings.txt", "r") as f:
        lines = f.readlines()
    lines = list(map(lambda s: s.strip(), lines))
    d = dict(x.split(":") for x in lines)
    return(d)

def getRating(ssValue):
    if  ssValue in range (0, 30):
        mastery = "Very Low"
        star    = "*"
    if  ssValue in range (30, 40):
        mastery = "Low"
        star    = "**"
    if  ssValue in range (40, 45):
        mastery = "Low Average"
        star    = "***"
    if  ssValue in range (45, 55):
        mastery = "Average"
        star    = "****"
    if  ssValue in range (55, 60):
        mastery = "High Average"
        star    = "*****"
    if  ssValue in range (60, 70):
        mastery = "High"
        star    = "******"
    if  ssValue >= 70:
        mastery = "Very High"
        star    = "*******"
    return mastery, star


def createDoc(dictOfSubjects, lengthOfSubjects, subheadingsDict, style):
    count = 0
    while count < lengthOfSubjects:
        #gets the name, email, and date completed for each subject
        name =  dictOfSubjects.get(count).get("Name")
        email =  dictOfSubjects.get(count).get("email")
        date  =  dictOfSubjects.get(count).get("Date").strftime("%m/%d/%y")
        
        word_document = Document()

        #adds the heading, name, email, and date completed to the document
        word_document.add_heading("#REDACTED")
        word_document.add_paragraph(f"Name: {name}\t Email: {email}\nDate completed: {date}")

        #adds the logo to the document in the header
        header = word_document.sections[0].header
        paragraph = header.paragraphs[0]
        logo_run = paragraph.add_run()
        logo_run.add_picture('fifthTheoryEducation.png', width=Inches(1.2))
        picture = word_document.sections[0].header.paragraphs[-1]
        picture.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        #adds the footer to the document
        footer = word_document.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.text = "Copyright © #REDACTED, LLC All rights reserved."

        document_name = name + " Report"

        #adds table to word doc and headings
        table = word_document.add_table(1,5)
        table.style = 'Table Grid'
        heading_cells = table.rows[0].cells
        heading_cells[0].text = "Assessment Sections"
        heading_cells[1].text = "Standard Score"
        heading_cells[2].text = "Percentile Score"
        heading_cells[3].text = "Overall Mastery"
        heading_cells[4].text = "7-Star Guide"

        for subject in dictOfSubjects.get(count).keys():
            # I chose to only deal and manipulate the SS values or Standard Score
            if ("SS" in subject):
                # Create a tempsubject so I can maniuplate it and leave the orginical in tact
                tempSubject = subject.replace("SS", "").strip()
                for heading, headingSubject in subheadingsDict.items():
                    #this for loop and if check for the case that the current SS value should have a subheading before it
                    #if it does, the subheading gets added and then the program continues normally
                    if tempSubject in headingSubject:
                        cells = table.add_row().cells
                        cells[0].text = heading
                        cells[0].paragraphs[0].runs[0].font.bold = True
                        cells[0].paragraphs[0].runs[0].font.size = Pt(10)

                #adding the values to the table. I add tempSubject as the value because it is already cleaned 
                #and devoid of SS
                cells = table.add_row().cells
                cells[0].text = tempSubject
                cells[0].paragraphs[0].runs[0].font.bold = False
                cells[0].paragraphs[0].runs[0].font.size = Pt(9)

                #get the score and write it to the table
                cells[1].text = str(dictOfSubjects.get(count).get(subject.strip()))
                cells[1].paragraphs[0].runs[0].font.bold = False
                cells[1].paragraphs[0].runs[0].font.size = Pt(9)

                #save the score so that I can determine the star rating / mastery
                ssValue = dictOfSubjects.get(count).get(subject)

                #search the dictionary for the coresponding PS value
                PSsubject = "PS " + tempSubject
                
                #saves the PS score to the table
                cells[2].text = str(dictOfSubjects.get(count).get(PSsubject.strip()))
                cells[2].paragraphs[0].runs[0].font.bold = False
                cells[2].paragraphs[0].runs[0].font.size = Pt(9)
               
                #adds the mastery and star guide
                mastery, star = getRating(ssValue)
                cells[3].text = mastery
                cells[4].text = star
                cells[3].paragraphs[0].runs[0].font.bold = False
                cells[4].paragraphs[0].runs[0].font.bold = False
                cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                cells[4].paragraphs[0].runs[0].font.size = Pt(9)

        word_document.save(style+".docx")
        count += 1

def convertAndDelete():
    directory = os.getcwd()
    convert(str(directory), str(directory) + "/dataPDF")
    files_in_directory = os.listdir(directory)
    docxFiles = [file for file in files_in_directory if file.endswith(".docx")]
    
    for file in docxFiles:
        path_to_file = os.path.join(directory, file)
        os.remove(path_to_file)



def main():
    dictOfSubjects, lengthOfSubjects = getDictionary()
    subheadingsDict                  = getSubHeadings()
    createDoc(dictOfSubjects, lengthOfSubjects, subheadingsDict, style)
    convertAndDelete()

main()

